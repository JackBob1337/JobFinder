from docx import Document
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn
from pathlib import Path
from typing import Iterable
from copy import deepcopy
from app.schemas.cv_content import CVEntry

SECTION_HEADERS = [
    "PROFESSIONAL SUMMARY",
    "TECHNICAL SKILLS",
    "PROFESSIONAL EXPERIENCE",
    "PROJECTS",
    "EDUCATION",
]

class CVParser: 
    def __init__(self, path: Path = Path("data/CV_Vialkov_Fullstack.docx")):
        self.path = path
        self.doc = Document(str(path))
        self.sections: dict[str, list[Paragraph]] = {}
        self.section_headers: list[str] = SECTION_HEADERS.copy()
        self.output_path: Path = path.with_name(f'{path.stem}_update{path.suffix}')

        self._parse_sections()

    def _header_detector(self, para: Paragraph) -> bool:
        text = para.text.strip()
        return any(run.bold for run in para.runs) and text.upper() in self.section_headers     

    def _check_name_in_sections(self, section_name: str) -> str:
        normalized_name = section_name.upper()

        if normalized_name not in SECTION_HEADERS:
            raise ValueError("Section doesn't exist")

        return normalized_name

    def _is_bullets(self, para: Paragraph) -> bool:
        return (
            para._element.pPr is not None
            and para._element.pPr.numPr is not None
        )    

    def _get_bullet_paragraphs(self, paragraphs: Iterable[Paragraph]) -> list[Paragraph]:
        return [
            para for para in paragraphs
            if self._is_bullets(para)
        ]

    def _parse_sections(self) -> dict[str, list[Paragraph]]:
        self.sections = {}
        current_section: str | None = None
        for para in self.doc.paragraphs:
            text = para.text.strip()

            if not text:
                continue

            if self._header_detector(para):
                current_section = text.upper()
                self.sections[current_section] = []
            elif current_section:
                self.sections[current_section].append(para)

        return self.sections
    
    def _update_bullet_text(self, bullet_paragraphs: list[Paragraph], bullets: list[str]) -> None:
        for para, text in zip(bullet_paragraphs, bullets):
            if para.runs:
                para.runs[0].text = text.strip()
                for run in para.runs[1:]:
                    run.text = ''

    def _resize_bullets(self, bullet_paragraphs: list[Paragraph], target_size: int) -> None:
        if len(bullet_paragraphs) < target_size:
            template = bullet_paragraphs[-1]

            for _ in range(target_size - len(bullet_paragraphs)):
                new_para = deepcopy(template._p)
                template._p.addnext(new_para)

                new_paragraph = Paragraph(new_para, template._parent)
                bullet_paragraphs.append(new_paragraph)

                template = new_paragraph

        elif len(bullet_paragraphs) > target_size:
            for _ in range(len(bullet_paragraphs) - target_size):
                para = bullet_paragraphs.pop()
                para._element.getparent().remove(para._element)

    def _set_paragraph_text(self, paragraph: Paragraph, text: str) -> None:
        if paragraph.runs:
            paragraph.runs[0].text = text.strip()

            for run in paragraph.runs[1:]:
                run.text = ''

        else:
            paragraph.add_run(text.strip())        

    def _clear_paragraph(self, paragraph: Paragraph) -> None:
        for run in paragraph.runs:
            run.text = ''
    
    def load_cv_text(self) -> str:
        return '\n'.join(para.text for para in self.doc.paragraphs if para.text.strip())
    
    def get_summary_text(self) -> str:
        normalized_name = "PROFESSIONAL SUMMARY"

        if normalized_name not in self.sections:
            raise ValueError('Section doesn`t exist')

        section_paragraphs = self.sections[normalized_name]
        lines = []
        for para in section_paragraphs:
            text = para.text.strip()
            if text:
                lines.append(text)
                
        return "\n".join(lines)

    def rewrite_summary(self, new_summary: str, section_name: str = 'PROFESSIONAL SUMMARY') -> None:
        normalized_name = self._check_name_in_sections(section_name)

        section_paragraphs = self.sections[normalized_name]

        if not section_paragraphs:
            raise ValueError("Section has no paragraphs to rewrite")

        self._set_paragraph_text(section_paragraphs[0], new_summary)

        for para in section_paragraphs[1:]:
            self._clear_paragraph(para)

        self._parse_sections()

    def save(self, output_path: Path | None = None) -> None:
        save_path = output_path or self.output_path
        self.doc.save(str(save_path))

    def list_sections(self) -> list[str]:
        return list(self.sections.keys())

        
    def get_bullets(self, section_name: str) -> list[str]:
        normalized_name = self._check_name_in_sections(section_name)

        paragraphs = self._get_bullet_paragraphs(self.sections[normalized_name])

        return [
            para.text.strip()
            for para in paragraphs
            if para.text.strip()
        ]

    def rewrite_bullets(self, bullets: list[str], section_name: str) -> None:
        normalized_name = self._check_name_in_sections(section_name)

        bullet_paragraphs = self._get_bullet_paragraphs(self.sections[normalized_name])

        if not bullet_paragraphs:
            raise ValueError("No bullet paragraphs found")

        self._resize_bullets(bullet_paragraphs, len(bullets))
        self._update_bullet_text(bullet_paragraphs, bullets)
        self._parse_sections()

    def _is_hyperlink(self, element) -> bool:
        return element.tag == qn('w:hyperlink')

    def _get_hyperlinks(self, paragraph: Paragraph) -> list[str]:
        hyperlinks = []

        for hyperlink in paragraph._p.findall(qn("w:hyperlink")):
            text = ''.join(
                node.text or ''
                for node in hyperlink.iter()
                if node.tag == qn("w:t")
            )

            if text.strip():
                hyperlinks.append(text.strip())

        return hyperlinks

    def get_non_hyperlink_text(self, paragraph: Paragraph) -> str:
        parts = []

        for element in paragraph._p:
            if self._is_hyperlink(element):
                continue

            for node in element.iter():
                if node.tag == qn('w:t') and node.text:
                    parts.append(node.text)

        return ''.join(parts)

    def _parse_stack(self, text: str) -> list[str]:
        stack_text = text.split(":", 1)[1]
        stack_text = stack_text.replace("·", "")

        return [
            item.strip()
            for item in stack_text.split(",")
            if item.strip()
        ]

    def _parse_entries(self, section_name: str) -> list[CVEntry]:
        normalized_name = self._check_name_in_sections(section_name)
        paragraphs = self.sections[normalized_name]

        entries: list[CVEntry] = []
        current: CVEntry | None = None
        
        for para in paragraphs:

            text = self.get_non_hyperlink_text(para).strip()

            if not text:
                continue

            if self._is_bullets(para):
                if current is not None:
                    current.bullets.append(text)
                continue

            if text.lower().startswith("stack:"):
                if current is not None:
                    current.stack = self._parse_stack(text)
                continue

            if current is not None:
                entries.append(current)

            current = CVEntry(
                title=text,
                stack=[],
                bullets=[],
            )

        if current is not None:
            entries.append(current)

        return entries

    def get_experience(self) -> list[CVEntry]:
        return self._parse_entries('PROFESSIONAL EXPERIENCE')

    def get_projects(self) -> list[CVEntry]:
        return self._parse_entries("PROJECTS")

